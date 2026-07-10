"""Resume text extraction + Groq-powered CV parsing.

Extraction path:
- PDF  -> Mistral OCR 3, merged with pypdf text when both are available
- DOCX -> python-docx (paragraphs + tables)

Contact fields (name, email, phone) are regex-extracted before the LLM call
and used to fill gaps the model misses — never to invent data not in the text.

Groq structures the full CV. The result is returned for student confirmation.
"""

from __future__ import annotations

import base64
import io
import re
from functools import lru_cache

import httpx
from docx import Document
from langchain_groq import ChatGroq
from pydantic import BaseModel, Field
from pypdf import PdfReader

from app.config import get_settings
from app.services.llm_utils import call_structured

MAX_CV_CHARS = 12000
MISTRAL_OCR_URL = "https://api.mistral.ai/v1/ocr"

EMAIL_RE = re.compile(
    r"\b[A-Za-z0-9][A-Za-z0-9._%+\-]*@[A-Za-z0-9][A-Za-z0-9.\-]*\.[A-Za-z]{2,}\b"
)
# Nepal + generic international phone patterns
PHONE_RE = re.compile(
    r"(?:"
    r"\+?977[\s\-]?[9][6-9]\d{8}"  # Nepal mobile +977
    r"|\+?977[\s\-]?\d{2}[\s\-]?\d{6,8}"  # Nepal landline-ish
    r"|\b9[78]\d{8}\b"  # local Nepal mobile
    r"|\+?\d{1,3}[\s\-]?\(?\d{2,4}\)?[\s\-]?\d{3,4}[\s\-]?\d{3,4}\b"
    r")"
)

HEADER_SKIP = frozenset(
    {
        "curriculum vitae",
        "resume",
        "résumé",
        "cv",
        "profile",
        "personal details",
        "contact",
        "contact information",
    }
)


class ExperienceEntry(BaseModel):
    title: str = Field(description="Job title or role held.")
    company: str | None = Field(None, description="Employer or organization name.")
    start_date: str | None = Field(None, description="Start date as written on CV, e.g. 'Jan 2023'.")
    end_date: str | None = Field(
        None,
        description="End date as written on CV, or 'Present' if current role.",
    )
    description: str | None = Field(None, description="Brief summary of responsibilities and achievements.")


class EducationEntry(BaseModel):
    degree: str = Field(description="Degree or qualification, e.g. 'BCA', 'BSc Computer Science'.")
    institution: str | None = Field(None, description="School, college, or university name.")
    year: str | None = Field(None, description="Graduation year or date range as written on CV.")


class ParsedCV(BaseModel):
    full_name: str | None = Field(None, description="Candidate's full name from the CV header.")
    email: str | None = Field(None, description="Email address if present.")
    phone: str | None = Field(None, description="Phone number if present.")
    summary: str | None = Field(None, description="Professional summary or objective, if present.")
    years_of_experience: float | None = Field(
        None,
        description="Total years of professional experience inferred from work history.",
    )
    education: list[EducationEntry] = Field(
        default_factory=list,
        description="Education entries, most recent first.",
    )
    experience: list[ExperienceEntry] = Field(
        default_factory=list,
        description="Work experience entries, most recent first.",
    )
    skills: list[str] = Field(
        default_factory=list,
        description="Concrete skills the candidate demonstrably has. No duplicates, no invented skills.",
    )
    suggested_target_role: str | None = Field(
        None,
        description="The job role this CV is most obviously aimed at, if clear.",
    )


class ContactHints(BaseModel):
    full_name: str | None = None
    email: str | None = None
    phone: str | None = None


def _normalize_phone(raw: str) -> str:
    cleaned = re.sub(r"\s+", " ", raw.strip())
    return cleaned


def _extract_email(text: str) -> str | None:
    match = EMAIL_RE.search(text)
    return match.group(0).lower() if match else None


def _extract_phone(text: str) -> str | None:
    for match in PHONE_RE.finditer(text):
        candidate = _normalize_phone(match.group(0))
        digits = re.sub(r"\D", "", candidate)
        if len(digits) >= 9:
            return candidate
    return None


def _guess_name_from_text(text: str) -> str | None:
    """Heuristic: first plausible name line near the top of the CV."""
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    for line in lines[:12]:
        lowered = line.lower().strip(":-•|")
        if lowered in HEADER_SKIP:
            continue
        if EMAIL_RE.search(line) or PHONE_RE.search(line):
            continue
        if re.search(r"https?://|linkedin\.com|github\.com|@/", line, re.I):
            continue
        if len(line) < 4 or len(line) > 55:
            continue
        if re.search(r"\d{4}", line):  # likely a date line
            continue
        words = line.split()
        if not 2 <= len(words) <= 5:
            continue
        if not all(re.match(r"^[A-Za-z][A-Za-z.'\-]*$", w) for w in words):
            continue
        # ALL CAPS names -> title case
        if line.isupper():
            return line.title()
        return line
    return None


def extract_contact_hints(cv_text: str) -> ContactHints:
    """Regex pre-pass for contact fields — only what appears in raw text."""
    return ContactHints(
        full_name=_guess_name_from_text(cv_text),
        email=_extract_email(cv_text),
        phone=_extract_phone(cv_text),
    )


def _merge_pdf_texts(parts: list[str]) -> str:
    """Merge OCR + pypdf blocks, deduplicating identical paragraphs."""
    seen: set[str] = set()
    blocks: list[str] = []
    for part in parts:
        for block in re.split(r"\n{2,}", part):
            block = block.strip()
            if not block or block in seen:
                continue
            seen.add(block)
            blocks.append(block)
    return "\n\n".join(blocks)


async def extract_text(filename: str, content: bytes) -> tuple[str, str]:
    """Return (text, extraction_method). Never hard-fails on OCR problems."""
    lowered = filename.lower()
    if lowered.endswith(".pdf"):
        parts: list[str] = []
        methods: list[str] = []
        settings = get_settings()

        if settings.mistral_api_key:
            try:
                ocr_text = await _ocr_pdf_with_mistral(content)
                if ocr_text.strip():
                    parts.append(ocr_text)
                    methods.append("mistral-ocr")
            except Exception as exc:
                print(f"Mistral OCR failed ({type(exc).__name__}: {exc}); using pypdf fallback")

        pypdf_text = _extract_pdf_text(content)
        if pypdf_text.strip():
            parts.append(pypdf_text)
            methods.append("pypdf")

        if not parts:
            return "", "none"

        merged = _merge_pdf_texts(parts)
        method = "+".join(methods) if len(methods) > 1 else methods[0]
        return merged, method

    if lowered.endswith(".docx"):
        return _extract_docx_text(content), "docx"
    raise ValueError("Unsupported file type — upload a .pdf or .docx resume.")


async def _ocr_pdf_with_mistral(content: bytes) -> str:
    settings = get_settings()
    encoded = base64.standard_b64encode(content).decode("ascii")
    async with httpx.AsyncClient(timeout=90.0) as client:
        response = await client.post(
            MISTRAL_OCR_URL,
            headers={"Authorization": f"Bearer {settings.mistral_api_key}"},
            json={
                "model": settings.mistral_ocr_model,
                "document": {
                    "type": "document_url",
                    "document_url": f"data:application/pdf;base64,{encoded}",
                },
            },
        )
        response.raise_for_status()
        pages = response.json().get("pages") or []
    return "\n\n".join(page.get("markdown") or "" for page in pages)


def _extract_pdf_text(content: bytes) -> str:
    reader = PdfReader(io.BytesIO(content))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx_text(content: bytes) -> str:
    document = Document(io.BytesIO(content))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells if cell.text.strip())
    return "\n".join(parts)


def _merge_hints(parsed: ParsedCV, hints: ContactHints) -> ParsedCV:
    """Fill only empty LLM fields from regex hints — never overwrite LLM values."""
    data = parsed.model_dump()
    if not data.get("full_name") and hints.full_name:
        data["full_name"] = hints.full_name
    if not data.get("email") and hints.email:
        data["email"] = hints.email
    if not data.get("phone") and hints.phone:
        data["phone"] = hints.phone
    return ParsedCV(**data)


def _dedupe_skills(skills: list[str]) -> list[str]:
    seen: set[str] = set()
    out: list[str] = []
    for skill in skills:
        normalized = skill.strip()
        if not normalized:
            continue
        key = normalized.lower()
        if key in seen:
            continue
        seen.add(key)
        out.append(normalized)
    return out


def build_parse_warnings(parsed: ParsedCV, hints: ContactHints) -> list[str]:
    warnings: list[str] = []
    if not parsed.full_name:
        warnings.append("Could not detect your full name — please enter it manually.")
    if not parsed.email:
        warnings.append("No email found in the CV — add it if you want it on your profile.")
    if not parsed.skills:
        warnings.append("No skills extracted — add skills manually before continuing.")
    if not parsed.education:
        warnings.append("No education entries found — you can add them below.")
    if hints.email and parsed.email and hints.email != parsed.email.lower():
        warnings.append("Email was verified from the raw CV text.")
    return warnings


@lru_cache
def _llm() -> ChatGroq:
    settings = get_settings()
    return ChatGroq(
        model=settings.groq_model,
        temperature=0.0,
        api_key=settings.groq_api_key,
    )


def _build_parse_prompt(cv_text: str, hints: ContactHints) -> str:
    hint_lines = []
    if hints.full_name:
        hint_lines.append(f"- Detected name (from text): {hints.full_name}")
    if hints.email:
        hint_lines.append(f"- Detected email (from text): {hints.email}")
    if hints.phone:
        hint_lines.append(f"- Detected phone (from text): {hints.phone}")
    hint_block = "\n".join(hint_lines) if hint_lines else "- No contact hints detected."

    return f"""You are a careful CV parser for DISHA AI, a Nepali career platform.

RULES (strict):
1. Extract ONLY information explicitly present in the resume text.
2. If a field is missing from the CV, leave it null or an empty list — NEVER invent.
3. Use the regex hints below when they match the CV; do not contradict the CV text.
4. Normalize skill names (ReactJS -> React, Node -> Node.js, PG -> PostgreSQL).
5. List education and experience most recent first.
6. years_of_experience: sum professional work durations if dates are clear; else null.
7. suggested_target_role: infer only if the CV clearly targets one role (e.g. titles, summary).

EXAMPLE output shape (illustrative — do not copy unless in the CV):
{{
  "full_name": "Sita Gurung",
  "email": "sita.gurung@email.com",
  "phone": "+977 9841112233",
  "years_of_experience": 2.0,
  "education": [{{"degree": "BSc CSIT", "institution": "Tribhuvan University", "year": "2023"}}],
  "experience": [{{"title": "Junior Developer", "company": "Leapfrog Technology", "start_date": "2023", "end_date": "Present", "description": "Built React dashboards"}}],
  "skills": ["Python", "React", "SQL", "Git"],
  "suggested_target_role": "Frontend Developer"
}}

REGEX HINTS (pre-extracted from raw text):
{hint_block}

RESUME TEXT:
{cv_text[:MAX_CV_CHARS]}
"""


async def parse_cv(cv_text: str) -> tuple[ParsedCV, list[str]]:
    hints = extract_contact_hints(cv_text)
    prompt = _build_parse_prompt(cv_text, hints)
    llm_result = await call_structured(_llm(), ParsedCV, prompt)

    if llm_result is None:
        # Regex-only fallback — contact fields only, no invented skills
        parsed = ParsedCV(
            full_name=hints.full_name,
            email=hints.email,
            phone=hints.phone,
        )
        warnings = build_parse_warnings(parsed, hints)
        warnings.insert(0, "AI parsing failed — contact fields were pre-filled where possible. Complete the rest manually.")
        return parsed, warnings

    parsed = _merge_hints(llm_result, hints)
    parsed.skills = _dedupe_skills(parsed.skills)
    warnings = build_parse_warnings(parsed, hints)
    return parsed, warnings


# Backward-compatible alias used by older imports.
parse_cv_skills = parse_cv
